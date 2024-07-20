from flask import Flask, render_template, request, send_file
import fitz  # PyMuPDF
from docx import Document
from io import BytesIO
import os

app = Flask(__name__)

def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(text)
    if style:
        p.style = style
    return p

def add_image(doc, image_stream, width=None):
    image_stream.seek(0)
    doc.add_picture(image_stream, width=width)

def convert_pdf_to_word(pdf_stream):
    doc = Document()
    pdf_document = fitz.open(stream=pdf_stream.read(), filetype="pdf")

    for page_number in range(len(pdf_document)):
        page = pdf_document.load_page(page_number)
        text = page.get_text()
        if text:
            add_paragraph(doc, text)

        images = page.get_images(full=True)
        for img_index, img in enumerate(images):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]
            image_stream = BytesIO(image_bytes)
            add_image(doc, image_stream)

    output_stream = BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        file = request.files['pdf']
        if file and file.filename.endswith('.pdf'):
            pdf_stream = file.stream
            word_stream = convert_pdf_to_word(pdf_stream)
            return send_file(word_stream, as_attachment=True, download_name='output.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        else:
            return "Invalid file format. Please upload a PDF file.", 400
    return render_template('index.html')

if __name__ == "__main__":
    app.run(debug=True)
